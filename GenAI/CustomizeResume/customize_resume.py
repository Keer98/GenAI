import docx
from docx import Document
from docx.shared import Pt
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain, SequentialChain
from langchain_core.output_parsers import JsonOutputParser
from resume_formatting import format_skills, format_responsibilities

def customize_resume(resume_file, job_description, llm):
    # Load the candidate's standard resume
    doc = docx.Document(resume_file)
    standard_resume = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    # Step 1: Define the first prompt template
    first_prompt = PromptTemplate(
        input_variables=["job_description", "resume"],
        template=""" You are a helpful assistant that evaluates how well a candidate's resume matches a given job description. You will carefully compare the two and provide a match score out of 100, along with a breakdown of the criteria used to compute the score. Use the following scoring criteria:
Job Description:
{job_description}
Candidate's Resume:
{resume}
Skills and Qualifications (40 points): How well the candidate's skills and qualifications match those listed in the job description.
Experience and Responsibilities (30 points): How closely the candidate's past experiences align with the job's responsibilities.
Keywords and Technical Terms (20 points): How many important keywords or technical terms from the job description are reflected in the resume.
Formatting and Relevance (10 points): Whether the resume is well-organized, relevant to the job, and free of unnecessary information.
Provide:

Match Score for Each Category: Show only the score for each category and the total score out of 100.
Key Points That Match: Highlight key aspects where the resume strongly aligns with the given job description.
Key Points That Don't Match: Highlight key aspects where the resume doesn't align with the given technical job description, skills, and responsibilities. (Don't include location). """
    )

    # Step 2: Define the second prompt template
    second_prompt = PromptTemplate(
        input_variables=["score_points", "job_description", "resume"],
        template="""I want to adjust my resume based on the job description and the points provided. 
points:{score_points}
job description: {job_description}
resume: {resume}
### Instructions:
A. **Update the Resume**: 
Adjust the resume based on the provided points. Currently try to add Key Points that Don't Match with job description from "points". 

Make changes to skills, experience, and summary sections by:
   - Adding missing skills to the appropriate subfields in the skills section.
   - Incorporating responsibilities into relevant experience sections.
   - Reflecting significant updates in the summary section.
B.  **Evaluate Updated Resume**: After updating, evaluate how well the new resume matches a given job description. You will carefully compare the two and provide a match score out of 100, along with a breakdown of the criteria used to compute the score. Use the following scoring criteria: :
  Skills and Qualifications (40 points): How well the candidate's skills and qualifications match those listed in the job description.
Experience and Responsibilities (30 points): How closely the candidate's past experiences align with the job's responsibilities.
Keywords and Technical Terms (20 points): How many important keywords or technical terms from the job description are reflected in the resume.
Formatting and Relevance (10 points): Whether the resume is well-organized, relevant to the job, and free of unnecessary information.
Provide:
Match Score for Each Category: Show only the score for each category and the total score out of 100.
C. **Newly Added Points**: Provide a list of key points added now to the resume during the update process based on the mismatches.
The overall goal is get the overall score near to 100 by incorporating the key points of job description that Don't Match with the given resume.
For Job Experience field you just give the list of Responsibilities no other fields
### JSON Output Keys:
- `Summary`
- `Skills`
- `1st Job Experience` 
- `2nd Job Experience` 
- `3rd Job Experience` 
- `5th Job Experience`
- `6th Job Experience`
- `7th Job Experience`
- `8th Job Experience`
- `9th Job Experience`
- `10th Job Experience` 
- 'Match Score'
-'Newly Added Points'

### VALID JSON FORMAT (NO PREAMBLE)
Only return the updated resume in JSON with the mentioned keys.
"""
    )

    # Step 3: Create individual chains
    first_chain = LLMChain(llm=llm, prompt=first_prompt, output_key="score_points")
    second_chain = LLMChain(llm=llm, prompt=second_prompt, output_key="combined_result")

    # Step 4: Combine chains into a Sequential Chain
    sequential_chain = SequentialChain(
        chains=[first_chain, second_chain],
        input_variables=["job_description", "resume"],  # Inputs for the first chain
        output_variables=["score_points", "combined_result"],  # Final outputs
    )

    # Step 5: Run the Chained Process
    output = sequential_chain({"job_description": job_description, "resume": standard_resume})
    score_points=output["score_points"]
    #customized_resume_points=output["customized_resume"]
    #print("Score Points:", output["score_points"])
    #print("Customized Resume:", output["customized_resume"])
    # Parse the JSON output
    json_parser = JsonOutputParser()
    json_cus_res = json_parser.parse(output["combined_result"])
    updated_score=json_cus_res["Match Score"]
    added_points=json_cus_res["Newly Added Points"]
    del json_cus_res["Match Score"]
    del json_cus_res["Newly Added Points"]
    json_cus_res["Skills"] = format_skills(json_cus_res["Skills"])
    
    for key in json_cus_res.keys():
        if 'Experience' in key:
            json_cus_res[key]=format_responsibilities(json_cus_res[key])
    #print(json_cus_res["Match Score for Each Category"])
    # Modify the Word document
    doc = Document("Formatted_Resume.docx")
    for paragraph in doc.paragraphs:
        for key, value in json_cus_res.items():
            if key in paragraph.text:
                # Convert value to a string based on its type
                if isinstance(value, dict):
                    # Convert dictionary to a formatted string
                    
                    value_str = "\n".join([f"{sub_key}: {sub_value}" for sub_key, sub_value in value.items()])
                elif isinstance(value, list):
                    # Convert list to a newline-separated string
                    value_str = "\n".join([str(item) for item in value])  # Ensure all items are strings
                else:
                    # If value is already a string, use it directly
                    value_str = value

                # Replace the placeholder with the formatted string
                paragraph.text = paragraph.text.replace(key, value_str)

                # Apply consistent formatting
                for run in paragraph.runs:
                    run.font.name = "Trebuchet MS"
                    run.font.size = Pt(11)
                    run.font.underline= False

    # Save the updated document
    output_file = "Updated_Resume.docx"
    doc.save(output_file)
    
    return output_file, score_points, updated_score,added_points
